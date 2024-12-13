from flask import render_template, request, jsonify
import subprocess
import threading
import time
from docx import Document
from datetime import datetime
from app import app  # Importamos la app para definir las rutas

# Ruta para mostrar el formulario o página inicial
@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")

# Ruta para manejar los comandos enviados desde el cliente
@app.route("/comando", methods=["POST"])
def ejecutar_comando():
    data = request.get_json()
    comando = data.get("comando")
    link = data.get("link")
    try:
        time_limit = int(data.get("temporizador", 30))  # Convertir a entero
    except ValueError:
        return jsonify({"error": "El temporizador debe ser un número válido."}), 400

    entry_time_limit = data.get("entry_time_limit")
    if entry_time_limit:
        try:
            entry_time_limit = int(entry_time_limit)  # Convertir a entero si se proporciona
        except ValueError:
            return jsonify({"error": "El tiempo de entrada debe ser un número válido."}), 400

    if not link:
        return jsonify({"error": "El enlace es obligatorio"}), 400

    if comando == "youtube":
        return open_video_on_devices(link, time_limit)
    elif comando == "enlace":
        selected_option = data.get("option", "default")
        return open_link_on_devices(link, selected_option, time_limit, entry_time_limit)
    else:
        return jsonify({"error": "Comando no reconocido"}), 400

# Función para registrar las acciones a modo de bitácora
def log_to_word(log_entries, filename="Bitacora.docx"):
    """
    Registra entradas en un archivo Word, agregando nuevas entradas a un archivo existente.
    log_entries: Lista de strings que contienen las entradas del log.
    filename: Nombre del archivo de Word.
    """
    try:
        # Intentar abrir el archivo existente
        try:
            doc = Document(filename)
        except Exception:
            # Si el archivo no existe, crear uno nuevo
            doc = Document()
            doc.add_heading('Bitácora de Proceso', level=1)
            doc.add_paragraph(f'Generado el: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph("")  # Línea vacía

        # Añadir nuevas entradas de log
        for entry in log_entries:
            doc.add_paragraph(entry)

        # Guardar el archivo actualizado
        doc.save(filename)
        print(f"Bitácora actualizada en {filename}")
    except Exception as e:
        print(f"Error al escribir la bitácora: {e}")


# Función para abrir enlaces generales y establecer el temporizador
def open_link_on_devices(link, selected_option, time_limit, entry_time_limit=None):
    log_entries = [f"Inicio del proceso para abrir el enlace: {link}"]
    try:
        devices_output = subprocess.check_output(["adb", "devices"]).decode("utf-8")
        devices = [line.split("\t")[0] for line in devices_output.splitlines() if "\tdevice" in line]

        if not devices:
            log_entries.append("No se encontraron dispositivos conectados.")
            log_to_word(log_entries)
            return {"error": "No se encontraron dispositivos conectados."}, 400

        for device in devices:
            if selected_option == "default":
                subprocess.run(["adb", "-s", device, "shell", "am", "start", "-a", "android.intent.action.VIEW", "-d", link])
            elif selected_option == "chrome":
                subprocess.run([
                    "adb", "-s", device, "shell", "am", "start", "-a", "android.intent.action.VIEW",
                    "-d", link, "-n", "com.android.chrome/com.google.android.apps.chrome.Main"
                ])
            elif selected_option == "opera":
                subprocess.run([
                    "adb", "-s", device, "shell", "am", "start", "-a", "android.intent.action.VIEW",
                    "-d", link, "com.opera.browser"
                ])
            elif selected_option == "firefox":
                subprocess.run([
                    "adb", "-s", device, "shell", "am", "start", "-a", "android.intent.action.VIEW",
                    "-d", link, "-n", "org.mozilla.firefox/org.mozilla.gecko.BrowserApp"
                ])
            elif selected_option == "brave":
                subprocess.run([
                    "adb", "-s", device, "shell", "am", "start", "-a", "android.intent.action.VIEW",
                    "-d", link, "-n", "com.brave.browser/org.chromium.chrome.browser.ChromeTabbedActivity"
                ])
            log_entries.append(f"Enlace abierto en el dispositivo: {device}")

        final_time_limit = int(entry_time_limit) if entry_time_limit else time_limit
        threading.Thread(target=start_timer_and_go_home, args=(final_time_limit, devices, log_entries)).start()

        log_entries.append(f"Temporizador de {final_time_limit} segundos iniciado para {len(devices)} dispositivos.")
        log_to_word(log_entries)
        return {"message": f"Enlace abierto en {len(devices)} dispositivos y el temporizador de {final_time_limit} segundos ha comenzado."}, 200
    except subprocess.CalledProcessError as e:
        log_entries.append(f"Hubo un problema ejecutando ADB: {e}")
        log_to_word(log_entries)
        return {"error": f"Hubo un problema ejecutando ADB: {e}"}, 500
    except Exception as e:
        log_entries.append(f"Ocurrió un error inesperado: {e}")
        log_to_word(log_entries)
        return {"error": f"Ocurrió un error inesperado: {e}"}, 500

# Función para abrir enlaces de YouTube y establecer el temporizador
def open_video_on_devices(link, time_limit):
    """
    Abre un video de YouTube en los dispositivos y establece un temporizador.
    :param link: Enlace del video de YouTube.
    :param time_limit: Tiempo de espera en segundos.
    :return: Respuesta JSON con el resultado del proceso.
    """
    if "youtube.com" not in link and "youtu.be" not in link:
        return {"error": "Por favor ingresa un enlace válido de YouTube."}, 400

    refresh_stop_events = {}
    refresh_counters = {}
    log_entries = [f"Inicio del proceso para abrir el video: {link}"]

    try:
        devices_output = subprocess.check_output(["adb", "devices"]).decode("utf-8")
        devices = [line.split("\t")[0] for line in devices_output.splitlines() if "device" in line]

        if not devices:
            log_entries.append("No se encontraron dispositivos conectados.")
            log_to_word(log_entries)
            return {"error": "No se encontraron dispositivos conectados."}, 400

        for device in devices:
            try:
                subprocess.run([
                    "adb", "-s", device, "shell", "am", "start", "-a", "android.intent.action.VIEW",
                    "-d", link, "-n", "com.android.chrome/com.google.android.apps.chrome.Main"
                ])
                start_page_refresh(device, refresh_stop_events, refresh_counters)
                log_entries.append(f"Video abierto exitosamente en el dispositivo: {device}")
            except Exception as e:
                log_entries.append(f"Error al abrir el video en el dispositivo {device}: {e}")

        threading.Thread(target=start_timer_and_go_home, args=(time_limit, devices)).start()

        log_entries.append(f"Temporizador de {time_limit} segundos iniciado para {len(devices)} dispositivos.")
        log_to_word(log_entries)
        return {"message": f"Video de YouTube abierto en {len(devices)} dispositivos y el temporizador ha comenzado."}, 200
    except subprocess.CalledProcessError as e:
        log_entries.append(f"Hubo un problema ejecutando ADB: {e}")
        log_to_word(log_entries)
        return {"error": f"Hubo un problema ejecutando ADB: {e}"}, 500
    except Exception as e:
        log_entries.append(f"Ocurrió un error inesperado: {e}")
        log_to_word(log_entries)
        return {"error": f"Ocurrió un error inesperado: {e}"}, 500

  # Función para refrescar la página cada 14 segundos
def start_page_refresh(device, refresh_stop_events, refresh_counters):
    """
    Inicia un proceso de refresco de página para un dispositivo.
    :param device: Identificador del dispositivo.
    :param refresh_stop_events: Diccionario para gestionar eventos de parada por dispositivo.
    :param refresh_counters: Diccionario para contar los refrescos por dispositivo.
    """
    stop_event = threading.Event()
    refresh_stop_events[device] = stop_event  # Asociar el evento al dispositivo

    # Inicializar el contador del dispositivo si no existe
    if device not in refresh_counters:
        refresh_counters[device] = 0

    def refresh_page():
        while not stop_event.is_set():  # Revisar si se debe detener
            try:
                subprocess.run(["adb", "-s", device, "shell", "input", "keyevent", "KEYCODE_F5"])
                # Incrementar el contador
                refresh_counters[device] += 1
                time.sleep(14)  # Refrescar cada 14 segundos
            except Exception as e:
                print(f"Error al refrescar la página en el dispositivo {device}: {e}")

    refresh_thread = threading.Thread(target=refresh_page, daemon=True)
    refresh_thread.start()

def start_timer_and_go_home(time_limit, devices, refresh_stop_events):
    """
    Maneja el temporizador y detiene el refresco en los dispositivos.
    :param time_limit: Tiempo de espera en segundos.
    :param devices: Lista de dispositivos conectados.
    :param refresh_stop_events: Diccionario para gestionar eventos de parada por dispositivo.
    """
    refresh_counters = {}
    log_entries = [f"Iniciando temporizador de {time_limit} segundos."]
    print(f"Iniciando temporizador de {time_limit} segundos.")
    time.sleep(time_limit)

    for device in devices:
        try:
            # Detener el refresco de página para el dispositivo
            if device in refresh_stop_events:
                refresh_stop_events[device].set()  # Detener el bucle del refresco
                del refresh_stop_events[device]   # Limpiar el evento de la lista

            # Enviar el dispositivo a la pantalla principal
            subprocess.run(["adb", "-s", device, "shell", "input", "keyevent", "KEYCODE_HOME"])
            log_entries.append(f"Dispositivo {device} enviado a la pantalla de inicio.")
        except Exception as e:
            log_entries.append(f"Error al intentar enviar {device} a inicio: {e}")

    log_entries.append(f"Temporizador de {time_limit} segundos completado.")
    log_to_word(log_entries)

